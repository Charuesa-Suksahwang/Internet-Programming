from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "คู่มือการใช้งาน_Cooking_Start.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
GRAY = "666666"


def set_font(run, size=11, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Cooking Start | หน้า ")
    set_font(run, 9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_font(run, 11, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_font(run, 11)
    else:
        run = p.add_run(text)
        set_font(run, 11)
    return p


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_font(run, 11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, 11)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    run = p.add_run(title + " ")
    set_font(run, 11, bold=True, color=DARK_BLUE)
    run = p.add_run(body)
    set_font(run, 11)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Cooking Start")
    set_font(run, 30, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("คู่มือการใช้งานระบบจัดการสินค้าและสต็อก")
    set_font(run, 16, color=GRAY)

    info = doc.add_table(rows=3, cols=2)
    set_table_geometry(info, [2700, 6660])
    for row, (label, value) in zip(info.rows, [
        ("วัตถุประสงค์", "ใช้จัดการข้อมูลสินค้า สต็อก และหมวดหมู่สินค้า"),
        ("ผู้ใช้งาน", "ผู้ดูแลระบบ (Admin)"),
        ("ระบบ", "Cooking Start Inventory Management"),
    ]):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        a = row.cells[0].paragraphs[0].add_run(label)
        set_font(a, 11, bold=True, color=DARK_BLUE)
        b = row.cells[1].paragraphs[0].add_run(value)
        set_font(b, 11)
    doc.add_paragraph()
    add_note(doc, "หมายเหตุ:", "ระบบต้องเชื่อมต่อกับ Backend จึงจะสามารถล็อกอินและบันทึกข้อมูลสินค้าได้")

    doc.add_page_break()

    add_heading(doc, "1. การเข้าสู่ระบบ")
    add_body(doc, "เมื่อเปิดแอป ระบบจะแสดงหน้า Sign in เพื่อป้องกันไม่ให้ผู้ที่ยังไม่ได้รับสิทธิ์จัดการข้อมูลสินค้า")
    add_numbered(doc, [
        "กรอก Username ที่ผู้ดูแลระบบกำหนด",
        "กรอก Password",
        "กดปุ่ม Log in",
        "หากข้อมูลถูกต้อง ระบบจะเข้าสู่หน้า Dashboard",
    ])
    add_note(doc, "หากเข้าสู่ระบบไม่ได้:", "ตรวจสอบว่า Backend กำลังทำงานอยู่ และตรวจสอบ Username กับ Password อีกครั้ง")

    add_heading(doc, "2. หน้า Dashboard")
    add_body(doc, "Dashboard เป็นหน้าสรุปข้อมูลล่าสุดของสินค้าในระบบ โดยจะแสดงข้อมูลจากฐานข้อมูลจริง")
    add_bullets(doc, [
        "Products: จำนวนรายการสินค้าทั้งหมด",
        "Total Stock: จำนวนสินค้าคงเหลือรวม",
        "Low Stock: จำนวนสินค้าที่สต็อกต่ำ",
        "Categories: จำนวนหมวดหมู่สินค้า",
        "Stores: จำนวนร้านหรือสถานที่จัดเก็บ",
    ])

    add_heading(doc, "3. ดูและค้นหาสินค้า")
    add_body(doc, "เลือกเมนู Products ที่แถบด้านล่าง เพื่อเปิดรายการสินค้าทั้งหมด")
    add_numbered(doc, [
        "พิมพ์ชื่อสินค้าหรือชื่อหมวดหมู่ในช่อง Search products",
        "ระบบจะแสดงเฉพาะสินค้าที่ตรงกับคำค้นหา",
        "แต่ละรายการแสดงรูปภาพ ชื่อสินค้า จำนวนสต็อก หมวดหมู่ และสถานที่เก็บสินค้า",
    ])

    doc.add_page_break()

    add_heading(doc, "4. เพิ่มสินค้าใหม่")
    add_body(doc, "กดปุ่ม + Add Product จากหน้า Products หรือเลือกเมนู Add ที่แถบด้านล่าง")
    add_body(doc, "ข้อมูลที่ต้องกรอกให้ครบ (มีเครื่องหมาย *) ได้แก่")
    add_bullets(doc, [
        "Name - ชื่อสินค้า",
        "Category - หมวดหมู่สินค้า",
        "Price - ราคาสินค้า",
        "Item Code - รหัสสินค้า",
        "Stock - จำนวนสินค้าในสต็อก",
    ])
    add_body(doc, "ข้อมูลเพิ่มเติมที่กรอกได้ เช่น Description, Image URL, Location, Status, Brand, Sizes และ Order Name")
    add_numbered(doc, [
        "กรอกข้อมูลสินค้า",
        "ตรวจสอบจำนวน Stock ต้องเป็นจำนวนเต็มตั้งแต่ 0 ขึ้นไป",
        "กดปุ่ม Save Product",
        "ระบบส่งข้อมูลไป Backend และบันทึกลง MySQL",
        "เมื่อสำเร็จ ระบบกลับไปหน้า Products และแสดงรายการล่าสุด",
    ])

    add_heading(doc, "5. แก้ไขสินค้า")
    add_numbered(doc, [
        "ไปที่หน้า Products",
        "กดสัญลักษณ์ดินสอของสินค้าที่ต้องการแก้ไข",
        "ระบบนำข้อมูลเดิมมาแสดงในฟอร์ม",
        "แก้ไขข้อมูลตามต้องการ แล้วกด Save Product",
        "ระบบอัปเดตข้อมูลสินค้าเดิม โดยไม่สร้างรายการใหม่",
    ])

    add_heading(doc, "6. ลบสินค้า")
    add_numbered(doc, [
        "ไปที่หน้า Products",
        "กดสัญลักษณ์ถังขยะของสินค้าที่ต้องการลบ",
        "ระบบแสดงกล่องยืนยันการลบ",
        "กด ลบ เพื่อยืนยัน หรือกด ยกเลิก เพื่อกลับไปหน้าเดิม",
        "เมื่อลบสำเร็จ ระบบจะโหลดรายการสินค้าใหม่ทันที",
    ])
    add_note(doc, "ข้อควรระวัง:", "การลบสินค้าไม่สามารถย้อนกลับได้ ควรตรวจสอบชื่อสินค้าก่อนกดยืนยัน")

    add_heading(doc, "7. การติดตั้งระบบบนเครื่องอื่น")
    add_body(doc, "ส่วนนี้ใช้เมื่ออาจารย์ต้องการดาวน์โหลดโค้ดจาก GitHub และทดลองระบบด้วยฐานข้อมูลของตนเอง")
    add_numbered(doc, [
        "ดาวน์โหลดโค้ดจาก GitHub และ import ไฟล์ ip_std6730202602.sql ใน phpMyAdmin เพื่อสร้างตารางและข้อมูลสินค้า",
        "เข้าโฟลเดอร์ backend แล้วคัดลอกไฟล์ .env.example เป็นชื่อ .env",
        "กรอกข้อมูลฐานข้อมูลใน backend/.env และกำหนด ADMIN_USERNAME กับ ADMIN_PASSWORD สำหรับใช้ล็อกอิน",
        "ติดตั้ง dependencies ในโฟลเดอร์ backend ด้วย npm install แล้วเริ่ม Backend ด้วย npm start",
        "ที่โฟลเดอร์หลักของโปรเจกต์ คัดลอก .env.example เป็น .env แล้วกำหนด EXPO_PUBLIC_API_BASE_URL ให้ตรงกับ Backend เช่น http://localhost:3056/api",
        "ติดตั้ง dependencies ของแอปด้วย npm install แล้วเริ่มแอปด้วย npx expo start",
        "ล็อกอินด้วย Username และ Password ที่กำหนดไว้ใน backend/.env",
    ])
    add_note(doc, "เรื่องความปลอดภัย:", "ไฟล์ .env ไม่มีใน GitHub เพราะเก็บข้อมูลฐานข้อมูลและรหัสผ่าน ส่วนไฟล์ .env.example เป็นเพียงแบบฟอร์มตัวอย่างสำหรับตั้งค่าบนเครื่องใหม่")

    add_heading(doc, "8. สรุปการทำงานของระบบ")
    add_body(doc, "Frontend รับข้อมูลและแสดงผลให้ผู้ใช้ ส่วน Backend ตรวจสอบสิทธิ์และบันทึกข้อมูลลงฐานข้อมูล MySQL")
    flow = doc.add_table(rows=1, cols=3)
    set_table_geometry(flow, [3120, 3120, 3120])
    for cell, text in zip(flow.rows[0].cells, ["Frontend\nหน้าจอแอป", "Backend\nตรวจสอบและจัดการข้อมูล", "MySQL\nจัดเก็บข้อมูลสินค้า"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for idx, line in enumerate(text.split("\n")):
            run = p.add_run(line)
            set_font(run, 10.5, bold=(idx == 0), color=DARK_BLUE if idx == 0 else None)
            if idx == 0:
                p.add_run("\n")

    # Appendix: user-provided screenshots for demonstration and submission.
    screenshots = [
        ("ภาพที่ 1 หน้าเข้าสู่ระบบ", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 190421.png"),
        ("ภาพที่ 2 หน้า Dashboard สรุปข้อมูลสินค้าและสต็อก", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191015.png"),
        ("ภาพที่ 3 หน้ารายการสินค้าและปุ่มเพิ่มสินค้า", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191158.png"),
        ("ภาพที่ 4 หน้ารายการสินค้าหลังเพิ่มข้อมูล", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191008.png"),
        ("ภาพที่ 5 ฟอร์มเพิ่มสินค้า", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 190839.png"),
        ("ภาพที่ 6 ฟอร์มแก้ไขสินค้า (ข้อมูลหลัก)", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191120.png"),
        ("ภาพที่ 7 ฟอร์มแก้ไขสินค้า (ข้อมูลเพิ่มเติมและปุ่มอัปเดต)", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191128.png"),
        ("ภาพที่ 8 หน้าต่างยืนยันก่อนลบสินค้า", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191144.png"),
        ("ภาพที่ 9 หน้าหมวดหมู่สินค้า", r"C:\Users\my\Pictures\Screenshots\Screenshot 2026-09-03 191053.png"),
    ]
    doc.add_page_break()
    add_heading(doc, "ภาพประกอบการใช้งาน")
    add_body(doc, "ภาพหน้าจอด้านล่างใช้ประกอบการสาธิตขั้นตอนการใช้งานระบบ")
    for index, (caption, image_path) in enumerate(screenshots):
        if index:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_font(run, 12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_after = Pt(10)
        pic = doc.add_picture(image_path, width=Inches(3.85))
        pic_p = doc.paragraphs[-1]
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.paragraph_format.space_after = Pt(8)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        set_font(run, 9, color=GRAY)

    doc.core_properties.title = "คู่มือการใช้งาน Cooking Start"
    doc.core_properties.subject = "User guide"
    doc.core_properties.author = "Cooking Start"
    doc.save(OUT)
    print("User guide created")


if __name__ == "__main__":
    build()
